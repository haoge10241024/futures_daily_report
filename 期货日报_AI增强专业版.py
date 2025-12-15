import os
import akshare as ak
import pandas as pd
from datetime import datetime, timedelta
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH
import matplotlib.pyplot as plt
from matplotlib import rcParams
import mplfinance as mpf
import streamlit as st
import requests
import json
import time
from bs4 import BeautifulSoup
import feedparser
import re

# 配置字体（完全避免Linux字体错误）
import warnings
import logging
import platform

# 彻底忽略所有matplotlib字体警告
warnings.filterwarnings('ignore')
logging.getLogger('matplotlib').setLevel(logging.CRITICAL)
logging.getLogger('matplotlib.font_manager').setLevel(logging.CRITICAL)

# 强制设置matplotlib使用默认字体，不使用中文字体
import matplotlib
matplotlib.rcParams['font.family'] = 'sans-serif'
matplotlib.rcParams['font.sans-serif'] = ['DejaVu Sans', 'Arial']
matplotlib.rcParams['axes.unicode_minus'] = False

# Windows本地开发时使用中文字体（仅用于兼容）
if platform.system() == 'Windows':
    try:
        matplotlib.rcParams['font.sans-serif'] = ['Microsoft YaHei', 'Arial']
    except:
        pass

# ============ API配置 ============
# 直接配置API密钥（内置密钥，无需用户配置）
DEFAULT_DEEPSEEK_API_KEY = "sk-293dec7fabb54606b4f8d4f606da3383"
DEFAULT_SERPER_API_KEY = "e62b7b8b688821905ad0d6c360e44813175792fc"
DEEPSEEK_API_URL = "https://api.deepseek.com/v1/chat/completions"

# 全局API密钥（直接使用，不需要用户配置）
DEEPSEEK_API_KEY = DEFAULT_DEEPSEEK_API_KEY
SERPER_API_KEY = DEFAULT_SERPER_API_KEY


class EnhancedNewsSearcher:
    """增强的期货新闻搜索器"""
    
    def __init__(self):
        self.session = requests.Session()
        self.session.headers.update({
            'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/91.0.4472.124 Safari/537.36'
        })
        
        # RSS订阅源
        self.rss_feeds = {
            '东方财富期货': 'http://feed.eastmoney.com/rss/futures.xml',
            '金融界期货': 'http://rss.jrj.com.cn/rss/futures_index.xml',
            '新浪财经': 'http://rss.sina.com.cn/finance/futures.xml'
        }
    
    def search_with_serper_api(self, commodity: str, days_back: int = 3, api_key: str = None, target_dates: list = None):
        """使用Serper API进行搜索（支持指定目标日期列表）"""
        if not api_key:
            return []
        
        try:
            # target_dates可以是单个日期字符串或日期列表
            if target_dates is None:
                dates_to_search = [datetime.now().strftime('%Y-%m-%d')]
            elif isinstance(target_dates, str):
                dates_to_search = [target_dates]
            elif isinstance(target_dates, list):
                dates_to_search = target_dates
            else:
                dates_to_search = [datetime.now().strftime('%Y-%m-%d')]
            
            all_news = []
            for target_date in dates_to_search:
                base_date = datetime.strptime(target_date, '%Y-%m-%d')
                
                # 搜索目标日期的新闻
                search_query = f'{commodity}期货 OR {commodity}价格 OR {commodity}市场 {base_date.strftime("%Y年%m月%d日")}'
                
                url = "https://google.serper.dev/search"
                payload = json.dumps({
                    "q": search_query,
                    "num": 50,  # 增加返回数量
                    "tbs": "qdr:d",  # 只搜索当天
                    "gl": "cn",
                    "hl": "zh-cn"
                })
                
                headers = {
                    'X-API-KEY': api_key,
                    'Content-Type': 'application/json'
                }
                
                response = requests.post(url, headers=headers, data=payload, timeout=30)
                
                if response.status_code == 200:
                    results = response.json()
                    
                    for item in results.get('organic', []):
                        if self._is_relevant_financial_news(item.get('title', ''), item.get('snippet', ''), commodity):
                            news_item = {
                                'title': item.get('title', ''),
                                'content': item.get('snippet', ''),
                                'url': item.get('link', ''),
                                'source': item.get('displayedLink', '财经资讯'),
                                'date': target_date,
                                'relevance': self._calculate_relevance(item.get('title', '') + item.get('snippet', ''), commodity)
                            }
                            all_news.append(news_item)
                
                time.sleep(0.5)  # 避免API限制
            
            # 返回所有搜索到的新闻（不限制数量）
            return all_news
                
        except Exception as e:
            print(f"  ❌ Serper搜索出错: {e}")
            return []
    
    def scrape_eastmoney_news(self, commodity: str, days_back: int = 3):
        """爬取东方财富期货新闻"""
        try:
            search_url = f"http://so.eastmoney.com/news/s?keyword={commodity}期货"
            
            response = self.session.get(search_url, timeout=15)
            response.encoding = 'utf-8'
            
            if response.status_code == 200:
                soup = BeautifulSoup(response.text, 'html.parser')
                news_list = []
                
                news_items = soup.find_all(['div', 'li'], class_=['news-item', 'result-item', 'item'])
                
                for item in news_items[:20]:
                    try:
                        title_elem = item.find(['a', 'h3'], href=True)
                        if title_elem:
                            title = title_elem.get_text(strip=True)
                            url = title_elem.get('href', '')
                            
                            content_elem = item.find(['p', 'div'], class_=['summary', 'content', 'desc'])
                            content = content_elem.get_text(strip=True) if content_elem else ''
                            
                            date_elem = item.find(['span', 'div'], class_=['time', 'date', 'publish-time'])
                            date_str = date_elem.get_text(strip=True) if date_elem else ''
                            
                            if title and self._is_relevant_financial_news(title, content, commodity):
                                news_item = {
                                    'title': title,
                                    'content': content,
                                    'url': url if url.startswith('http') else f"http://futures.eastmoney.com{url}",
                                    'source': '东方财富',
                                    'date': self._parse_date(date_str),
                                    'relevance': self._calculate_relevance(title + content, commodity)
                                }
                                news_list.append(news_item)
                    except Exception:
                        continue
                
                return sorted(news_list, key=lambda x: x['relevance'], reverse=True)[:10]
            else:
                return []
                
        except Exception as e:
            print(f"  ❌ 东方财富爬取出错: {e}")
            return []
    
    def scrape_jrj_news(self, commodity: str, days_back: int = 3):
        """爬取金融界期货新闻"""
        try:
            search_url = f"http://search.jrj.com.cn/?q={commodity}期货&t=news"
            
            response = self.session.get(search_url, timeout=15)
            response.encoding = 'gbk'
            
            if response.status_code == 200:
                soup = BeautifulSoup(response.text, 'html.parser')
                news_list = []
                
                news_items = soup.find_all(['div', 'li'], class_=['search-result', 'news-item'])
                
                for item in news_items[:15]:
                    try:
                        title_elem = item.find('a')
                        if title_elem:
                            title = title_elem.get_text(strip=True)
                            url = title_elem.get('href', '')
                            
                            content_elem = item.find(['p', 'span'], class_=['summary', 'desc'])
                            content = content_elem.get_text(strip=True) if content_elem else ''
                            
                            if title and self._is_relevant_financial_news(title, content, commodity):
                                news_item = {
                                    'title': title,
                                    'content': content,
                                    'url': url,
                                    'source': '金融界',
                                    'date': datetime.now().strftime('%Y-%m-%d'),
                                    'relevance': self._calculate_relevance(title + content, commodity)
                                }
                                news_list.append(news_item)
                    except Exception:
                        continue
                
                return sorted(news_list, key=lambda x: x['relevance'], reverse=True)[:10]
            else:
                return []
                
        except Exception as e:
            print(f"  ❌ 金融界爬取出错: {e}")
            return []
    
    def get_rss_news(self, commodity: str, days_back: int = 3):
        """从RSS订阅源获取新闻"""
        try:
            all_news = []
            cutoff_date = datetime.now() - timedelta(days=days_back)
            
            for feed_name, feed_url in self.rss_feeds.items():
                try:
                    feed = feedparser.parse(feed_url)
                    
                    for entry in feed.entries[:20]:
                        title = entry.get('title', '')
                        summary = entry.get('summary', entry.get('description', ''))
                        link = entry.get('link', '')
                        
                        pub_date = entry.get('published_parsed')
                        if pub_date:
                            pub_datetime = datetime(*pub_date[:6])
                            if pub_datetime < cutoff_date:
                                continue
                        
                        if self._is_relevant_financial_news(title, summary, commodity):
                            news_item = {
                                'title': title,
                                'content': summary,
                                'url': link,
                                'source': feed_name.replace('_RSS', ''),
                                'date': pub_datetime.strftime('%Y-%m-%d') if pub_date else datetime.now().strftime('%Y-%m-%d'),
                                'relevance': self._calculate_relevance(title + summary, commodity)
                            }
                            all_news.append(news_item)
                    
                    time.sleep(1)
                    
                except Exception as e:
                    continue
            
            return sorted(all_news, key=lambda x: x['relevance'], reverse=True)[:10]
            
        except Exception as e:
            return []
    
    def _is_relevant_financial_news(self, title: str, content: str, commodity: str) -> bool:
        """判断新闻是否与期货品种相关"""
        text = (title + ' ' + content).lower()
        
        futures_keywords = ['期货', '价格', '市场', '合约', '交易', '涨跌', '行情', '分析', '预测']
        commodity_keywords = [commodity.lower(), f'{commodity}价格', f'{commodity}市场']
        
        has_commodity = any(keyword in text for keyword in commodity_keywords)
        has_futures = any(keyword in text for keyword in futures_keywords)
        
        return has_commodity and has_futures
    
    def _calculate_relevance(self, text: str, commodity: str) -> float:
        """计算新闻与商品的相关性得分"""
        text = text.lower()
        score = 0.0
        
        if commodity.lower() in text:
            score += 5.0
        
        futures_words = ['期货', '价格', '涨跌', '行情', '合约', '交易']
        for word in futures_words:
            if word in text:
                score += 1.0
        
        timely_words = ['今日', '昨日', '最新', '最近', '今天']
        for word in timely_words:
            if word in text:
                score += 0.5
        
        return min(score, 10.0)
    
    def _parse_date(self, date_str: str) -> str:
        """解析日期字符串"""
        if not date_str:
            return datetime.now().strftime('%Y-%m-%d')
        
        try:
            date_str = re.sub(r'[年月]', '-', date_str)
            date_str = re.sub(r'日', '', date_str)
            
            if '今天' in date_str or '今日' in date_str:
                return datetime.now().strftime('%Y-%m-%d')
            elif '昨天' in date_str or '昨日' in date_str:
                return (datetime.now() - timedelta(days=1)).strftime('%Y-%m-%d')
            else:
                return datetime.now().strftime('%Y-%m-%d')
        except:
            return datetime.now().strftime('%Y-%m-%d')
    
    def search_professional_data(self, commodity: str, serper_key: str, target_date: str = None):
        """搜索专业期货分析数据（8大维度）"""
        try:
            if not target_date:
                target_date = datetime.now().strftime('%Y-%m-%d')
            
            # 8大专业分析维度的搜索关键词
            professional_keywords = {
                # 1. 库存仓单数据
                '库存仓单': [
                    f"{commodity} 仓单 {target_date}",
                    f"{commodity} 库存 {target_date}",
                    f"{commodity} 交易所仓单"
                ],
                # 2. 基差数据
                '基差分析': [
                    f"{commodity} 基差 {target_date}",
                    f"{commodity} 期现价差",
                    f"{commodity} 现货价格 期货价格"
                ],
                # 3. 期限结构
                '期限结构': [
                    f"{commodity} 月差 {target_date}",
                    f"{commodity} 远近月价差",
                    f"{commodity} 跨期价差"
                ],
                # 4. 持仓席位
                '持仓席位': [
                    f"{commodity} 持仓席位 {target_date}",
                    f"{commodity} 主力持仓",
                    f"{commodity} 多空持仓"
                ],
                # 5. 供需数据
                '供需分析': [
                    f"{commodity} 产量 {target_date}",
                    f"{commodity} 消费量",
                    f"{commodity} 供需平衡表"
                ],
                # 6. 产业链数据
                '产业链': [
                    f"{commodity} 产业链 价格",
                    f"{commodity} 上下游",
                    f"{commodity} 生产利润"
                ],
                # 7. 进出口数据
                '进出口': [
                    f"{commodity} 进口量 {target_date}",
                    f"{commodity} 出口量",
                    f"{commodity} 海关数据"
                ],
                # 8. 宏观政策
                '宏观政策': [
                    f"{commodity} 政策 {target_date}",
                    f"{commodity} 行业政策",
                    f"{commodity} 国家政策"
                ]
            }
            
            professional_data = {}
            url = "https://google.serper.dev/search"
            
            for category, keywords in professional_keywords.items():
                category_data = []
                
                for keyword in keywords[:2]:  # 每个维度搜索前2个关键词
                    try:
                        payload = json.dumps({
                            "q": keyword,
                            "num": 5,
                            "gl": "cn",
                            "hl": "zh-cn"
                        })
                        
                        headers = {
                            'X-API-KEY': serper_key,
                            'Content-Type': 'application/json'
                        }
                        
                        response = requests.post(url, headers=headers, data=payload, timeout=30)
                        
                        if response.status_code == 200:
                            results = response.json()
                            for item in results.get('organic', [])[:2]:  # 每个关键词取前2条
                                category_data.append({
                                    'title': item.get('title', ''),
                                    'content': item.get('snippet', ''),
                                    'url': item.get('link', ''),
                                    'source': item.get('displayedLink', '未知'),
                                    'date': target_date,
                                    'category': category
                                })
                        
                        time.sleep(0.5)  # 避免API频率限制
                    except Exception as e:
                        print(f"  ⚠️ 搜索{category}数据出错: {e}")
                        continue
                
                if category_data:
                    professional_data[category] = category_data
            
            return professional_data
        except Exception as e:
            print(f"  ⚠️ 专业数据搜索出错: {e}")
            return {}
    
    def comprehensive_search(self, commodity: str, days_back: int = 3, serper_key: str = None, target_dates: list = None):
        """综合搜索：结合多种数据源（支持日期列表）"""
        all_news = []
        
        # 1. Serper API搜索商品新闻（传递target_dates日期列表）
        if serper_key:
            serper_news = self.search_with_serper_api(commodity, days_back, serper_key, target_dates)
            all_news.extend(serper_news)
        
        # 2. 网页爬虫
        eastmoney_news = self.scrape_eastmoney_news(commodity, days_back)
        all_news.extend(eastmoney_news)
        
        jrj_news = self.scrape_jrj_news(commodity, days_back)
        all_news.extend(jrj_news)
        
        # 3. RSS订阅
        rss_news = self.get_rss_news(commodity, days_back)
        all_news.extend(rss_news)
        
        # 去重和排序
        seen_titles = set()
        unique_news = []
        
        for news in all_news:
            title = news['title']
            if title not in seen_titles:
                seen_titles.add(title)
                unique_news.append(news)
        
        unique_news.sort(key=lambda x: x['relevance'], reverse=True)
        
        return unique_news[:25]


# ============ 技术指标计算函数 ============

def calculate_technical_indicators(market_data_df):
    """计算技术指标（MA、MACD、RSI等）"""
    try:
        if market_data_df.empty or len(market_data_df) < 20:
            return {}
        
        close_prices = market_data_df['close']
        
        # 计算均线
        ma5 = close_prices.rolling(window=5).mean().iloc[-1] if len(close_prices) >= 5 else None
        ma10 = close_prices.rolling(window=10).mean().iloc[-1] if len(close_prices) >= 10 else None
        ma20 = close_prices.rolling(window=20).mean().iloc[-1] if len(close_prices) >= 20 else None
        
        # 计算MACD
        exp1 = close_prices.ewm(span=12, adjust=False).mean()
        exp2 = close_prices.ewm(span=26, adjust=False).mean()
        macd = exp1 - exp2
        signal = macd.ewm(span=9, adjust=False).mean()
        macd_hist = macd - signal
        
        # 计算RSI
        delta = close_prices.diff()
        gain = (delta.where(delta > 0, 0)).rolling(window=14).mean()
        loss = (-delta.where(delta < 0, 0)).rolling(window=14).mean()
        rs = gain / loss
        rsi = 100 - (100 / (1 + rs))
        
        # 计算布林带
        ma20_series = close_prices.rolling(window=20).mean()
        std20 = close_prices.rolling(window=20).std()
        upper_band = ma20_series + (std20 * 2)
        lower_band = ma20_series - (std20 * 2)
        
        current_price = close_prices.iloc[-1]
        
        return {
            'ma5': round(ma5, 2) if ma5 else None,
            'ma10': round(ma10, 2) if ma10 else None,
            'ma20': round(ma20, 2) if ma20 else None,
            'macd': round(macd.iloc[-1], 2) if len(macd) > 0 else None,
            'macd_signal': round(signal.iloc[-1], 2) if len(signal) > 0 else None,
            'macd_hist': round(macd_hist.iloc[-1], 2) if len(macd_hist) > 0 else None,
            'rsi': round(rsi.iloc[-1], 2) if len(rsi) > 0 and not pd.isna(rsi.iloc[-1]) else None,
            'upper_band': round(upper_band.iloc[-1], 2) if len(upper_band) > 0 else None,
            'lower_band': round(lower_band.iloc[-1], 2) if len(lower_band) > 0 else None,
            'current_price': round(current_price, 2),
            'price_position': '上轨附近' if current_price > upper_band.iloc[-1] else ('下轨附近' if current_price < lower_band.iloc[-1] else '中轨附近') if len(upper_band) > 0 else '中轨'
        }
    except Exception as e:
        print(f"  ⚠️ 技术指标计算出错: {e}")
        return {}


# ============ AI辅助函数 ============

def ai_generate_market_description(market_data: dict, commodity_name: str, date_str: str, technical_indicators: dict = None, market_date=None) -> str:
    """使用DeepSeek AI生成行情描述（强化真实性约束）
    
    Args:
        market_data: 市场数据字典
        commodity_name: 商品名称
        date_str: 日报生成日期
        technical_indicators: 技术指标
        market_date: 实际的市场回顾日期（交易日），如果为None则自动计算前一天
    """
    print(f"[INFO] ========== AI生成行情描述 ==========")
    print(f"[INFO] 品种: {commodity_name}, 日报日期: {date_str}")
    
    # 如果提供了实际的市场日期，直接使用；否则计算前一天
    if market_date is None:
        report_date = datetime.strptime(date_str, '%Y-%m-%d')
        market_date = (report_date - timedelta(days=1)).strftime('%Y-%m-%d')
    elif isinstance(market_date, datetime):
        market_date = market_date.strftime('%Y-%m-%d')
    
    print(f"[INFO] 市场回顾日期: {market_date}")
    print(f"[INFO] API密钥长度: {len(DEEPSEEK_API_KEY)}")
    print(f"[INFO] API URL: {DEEPSEEK_API_URL}")
    try:
        prompt = f"""
你是专业期货分析师，正在为{date_str}撰写{commodity_name}期货的市场走势描述。

【重要约束】
1. ⚠️ 必须严格基于下方提供的真实行情数据，严禁编造任何价格或涨跌幅
2. ⚠️ 如果某项数据显示为"N/A"，不要推测或编造，直接说明数据缺失
3. ⚠️ 引用的所有数字必须与提供的数据完全一致，不要四舍五入或修改
4. ⚠️ 分析日期为{date_str}，不要使用其他日期

【{commodity_name}期货真实行情数据 - {date_str}】

日盘数据：
- 开盘价：{market_data.get('open', 'N/A')}元
- 收盘价：{market_data.get('close', 'N/A')}元
- 最高价：{market_data.get('high', 'N/A')}元
- 最低价：{market_data.get('low', 'N/A')}元
- 涨跌额：{market_data.get('change', 'N/A')}元
- 涨跌幅：{market_data.get('change_pct', 'N/A')}%

夜盘数据：
- 开盘价：{market_data.get('night_open', 'N/A')}元
- 收盘价：{market_data.get('night_close', 'N/A')}元
- 涨跌额：{market_data.get('night_change', 'N/A')}元
- 涨跌幅：{market_data.get('night_change_pct', 'N/A')}%

【撰写要求】
请撰写一段160-200字的专业行情走势描述，必须包含：
1. 日盘走势 - 开盘→盘中高低点→收盘，描述价格运行轨迹
2. 夜盘走势 - 开盘→收盘，与日盘对比分析
3. 技术形态 - 如"十字星"、"长阳线"、"跳空"等（基于真实数据判断）
4. 关键点位 - 支撑位、压力位（基于给定的高低价）
5. 市场情绪 - 多空力量对比（基于涨跌推断）

【文风要求】
- 使用专业术语：如"承压回落"、"获得支撑"、"震荡整理"等
- 语言精准：每个价格都要有来源依据
- 逻辑清晰：日盘→夜盘→全天总结
- 客观描述：不添加未提供的成交量、持仓量等数据

⚠️ 再次强调：所有数字必须来源于上述真实数据，不要编造！

请直接输出描述文本（不要标题、不要前言）：
"""
        
        headers = {
            "Authorization": f"Bearer {DEEPSEEK_API_KEY}",
            "Content-Type": "application/json"
        }
        
        data = {
            "model": "deepseek-chat",
            "messages": [{"role": "user", "content": prompt}],
            "temperature": 0.2,  # 降低温度，确保数据准确性
            "max_tokens": 600
        }
        
        response = requests.post(DEEPSEEK_API_URL, headers=headers, json=data, timeout=60)
        
        print(f"[DEBUG] AI生成行情描述 - 状态码: {response.status_code}")
        
        if response.status_code == 200:
            result = response.json()
            content = result['choices'][0]['message']['content'].strip()
            print(f"[DEBUG] AI生成行情描述 - 返回内容长度: {len(content)}字符")
            print(f"[DEBUG] AI生成行情描述 - 前100字符: {content[:100]}")
            return content
        else:
            error_msg = f"AI生成失败 (状态码: {response.status_code})"
            print(f"[ERROR] {error_msg}")
            return error_msg
            
    except Exception as e:
        error_msg = f"AI生成出错: {str(e)}"
        print(f"[ERROR] {error_msg}")
        return error_msg


def ai_generate_news_summary(commodity_name: str, date_str: str, news_list: list, professional_data: dict = None) -> str:
    """使用DeepSeek AI生成新闻资讯摘要"""
    try:
        # 准备所有新闻数据
        all_news = []
        ref_index = 1
        
        # 商品新闻
        for news in news_list[:15]:  # 最多15条商品新闻
            all_news.append({
                'index': ref_index,
                'title': news.get('title', ''),
                'content': news.get('content', ''),
                'source': news.get('source', '未知'),
                'date': news.get('date', 'N/A'),
                'url': news.get('url', '无')
            })
            ref_index += 1
        
        # 专业维度数据
        if professional_data:
            for category, data_list in professional_data.items():
                for data in data_list[:2]:  # 每个维度2条
                    all_news.append({
                        'index': ref_index,
                        'title': f"[{category}] {data.get('title', '')}",
                        'content': data.get('content', ''),
                        'source': data.get('source', '未知'),
                        'date': date_str,
                        'url': data.get('url', '无')
                    })
                    ref_index += 1
        
        # 构建新闻摘要
        news_summary = ""
        for news in all_news:
            news_summary += f"[{news['index']}] {news['title']}\n"
            if news['content']:
                content = news['content'][:200].replace('...', '').replace('…', '')
                news_summary += f"    {content}\n"
            news_summary += f"    来源：{news['source']} | 日期：{news['date']}\n\n"
        
        prompt = f"""
你是专业期货分析师，正在为{date_str}的{commodity_name}期货日报整理新闻资讯。

【重要约束】
1. ⚠️ 必须基于下方提供的真实新闻，严禁编造任何信息
2. ⚠️ 只提取和输出新闻的主体内容，不要添加评论或分析
3. ⚠️ 每条新闻保持独立，不要合并或重写
4. ⚠️ 使用上标[1][2][3]标注新闻来源
5. ⚠️ 日期为{date_str}，优先选择当天或临近日期的新闻

【新闻原始数据】
{news_summary}

【整理要求】
请将上述**所有新闻**整理输出，要求：

1. **输出所有新闻** - 不要筛选或删除任何新闻，全部输出
2. **时效性排序** - 按照日期排序，{date_str}的新闻排在前面
3. **内容完整** - 保留新闻主体内容，不截断
4. **去除重复** - 相似内容只保留一条
5. **格式统一** - 每条新闻格式一致

【输出格式】
序号. 新闻主体内容[上标序号]

例如：
1. 今日生猪期货主力合约收于11325元，日内下跌2.12%，创近期新低[1]

2. 全国生猪均价跌至11.19元/公斤，环比下跌0.26元，月环比下跌16.04%[2]

3. 交易所仓单数据显示，生猪仓单环比增加3000吨至15000吨，供应压力持续[3]

...（输出所有新闻）

⚠️ 重要提示：
- 直接输出新闻内容，不要添加"资讯汇总"、"市场动态"等标题
- 不要对新闻进行评论、分析或总结
- 不要合并多条新闻
- 保持新闻的客观性和独立性
- 每条新闻末尾必须有上标引用[数字]
- **必须输出所有提供的新闻，不要遗漏**

请直接输出整理后的新闻列表：
"""
        
        headers = {
            "Authorization": f"Bearer {DEEPSEEK_API_KEY}",
            "Content-Type": "application/json"
        }
        
        data = {
            "model": "deepseek-chat",
            "messages": [{"role": "user", "content": prompt}],
            "temperature": 0.1,  # 极低温度，确保忠实原文
            "max_tokens": 4000   # 支持输出所有新闻（约20-30条）
        }
        
        response = requests.post(DEEPSEEK_API_URL, headers=headers, json=data, timeout=90)
        
        if response.status_code == 200:
            result = response.json()
            return result['choices'][0]['message']['content'].strip()
        else:
            return f"AI生成失败 (状态码: {response.status_code})"
            
    except Exception as e:
        return f"AI生成出错: {str(e)}"


def ai_generate_main_view(commodity_name: str, date_str: str, market_data: dict, news_list: list, 
                         professional_data: dict = None, technical_indicators: dict = None, market_date=None) -> str:
    """使用DeepSeek AI生成主要观点（专业版：8大维度分析）
    
    Args:
        commodity_name: 商品名称
        date_str: 日报生成日期
        market_data: 市场数据
        news_list: 新闻列表
        professional_data: 专业数据
        technical_indicators: 技术指标
        market_date: 实际的市场回顾日期（交易日），如果为None则自动计算前一天
    """
    try:
        # 如果提供了实际的市场日期，直接使用；否则计算前一天
        if market_date is None:
            report_date = datetime.strptime(date_str, '%Y-%m-%d')
            market_date = (report_date - timedelta(days=1)).strftime('%Y-%m-%d')
        elif isinstance(market_date, datetime):
            market_date = market_date.strftime('%Y-%m-%d')
        # 准备商品新闻摘要（使用上标引用格式）
        news_summary = ""
        ref_index = 1
        for i, news in enumerate(news_list[:10], 1):
            news_summary += f"[{ref_index}] {news['title']}\n"
            if news.get('content'):
                content = news['content'][:150].replace('...', '').replace('…', '')
                news_summary += f"    {content}\n"
            news_summary += f"    来源：{news.get('source', '未知')} | 日期：{news.get('date', 'N/A')} | URL：{news.get('url', '无')}\n"
            ref_index += 1
        
        if not news_summary:
            news_summary = "暂无最新商品新闻数据"
        
        # 准备技术指标摘要
        tech_summary = "\n【技术分析指标】\n"
        if technical_indicators and any(technical_indicators.values()):
            tech_summary += f"当前价格：{technical_indicators.get('current_price', 'N/A')}元\n"
            tech_summary += f"MA5：{technical_indicators.get('ma5', 'N/A')}元\n"
            tech_summary += f"MA10：{technical_indicators.get('ma10', 'N/A')}元\n"
            tech_summary += f"MA20：{technical_indicators.get('ma20', 'N/A')}元\n"
            tech_summary += f"MACD：{technical_indicators.get('macd', 'N/A')} (信号线：{technical_indicators.get('macd_signal', 'N/A')})\n"
            tech_summary += f"RSI(14)：{technical_indicators.get('rsi', 'N/A')}\n"
            tech_summary += f"布林带：上轨{technical_indicators.get('upper_band', 'N/A')}元 / 下轨{technical_indicators.get('lower_band', 'N/A')}元\n"
            tech_summary += f"价格位置：{technical_indicators.get('price_position', 'N/A')}\n"
        else:
            tech_summary += "技术指标数据不足\n"
        
        # 准备8大维度专业数据摘要
        professional_summary = "\n【专业维度数据】\n"
        if professional_data and len(professional_data) > 0:
            for category, data_list in professional_data.items():
                professional_summary += f"\n◆ {category}：\n"
                for data in data_list[:2]:  # 每个维度最多2条
                    professional_summary += f"[{ref_index}] {data.get('title', '')}\n"
                    if data.get('content'):
                        content = data['content'][:100].replace('...', '').replace('…', '')
                        professional_summary += f"    {content}\n"
                    professional_summary += f"    来源：{data.get('source', '未知')} | URL：{data.get('url', '无')}\n"
                    ref_index += 1
        else:
            professional_summary += "暂无专业维度数据\n"
        
        prompt = f"""
你是资深期货分析师，正在为{date_str}撰写{commodity_name}期货的专业市场分析报告。

【重要约束】
1. ⚠️ 所有分析必须基于下方提供的真实数据，严禁编造任何信息
2. ⚠️ 如果某项数据显示为"N/A"或"暂无"，不要推测或编造该数据
3. ⚠️ 引用的所有价格、涨跌幅、指标值必须与提供的数据完全一致
4. ⚠️ 报告日期为{date_str}，不要使用其他日期的信息
5. ⚠️ 必须基于8大分析维度进行系统性分析

【一、价格数据 - {date_str}】
日盘收盘价：{market_data.get('close', 'N/A')}元
日盘涨跌幅：{market_data.get('change_pct', 'N/A')}%
夜盘收盘价：{market_data.get('night_close', 'N/A')}元
夜盘涨跌幅：{market_data.get('night_change_pct', 'N/A')}%
最高价：{market_data.get('high', 'N/A')}元
最低价：{market_data.get('low', 'N/A')}元
{tech_summary}

【二、市场新闻资讯】
{news_summary}
{professional_summary}

【撰写要求 - 8大维度专业分析框架】⭐核心

请撰写一段400-500字的深度主要观点，必须包含以下结构和维度：

**结构框架：**
【早盘聚焦】（120字）
- 今日最新消息分析
- 开盘前突发事件影响评估
- 与昨日行情的关联分析

【深度分析】（220字）- ⭐核心，必须覆盖8大维度
1. **技术面**：均线系统、RSI信号、支撑压力位
2. **基本面**：供需格局、库存仓单、基差走势[引用数据]
3. **资金面**：持仓席位、主力动向[引用数据]
4. **产业链**：上下游价格、利润[引用数据]
5. **政策面**：最新政策影响[引用数据]
6. **进出口**：海关数据[引用数据]
7. **市场情绪**：新闻舆情变化[引用数据]
8. **风险因素**：关键风险点[引用数据]

【今日观点及操作建议】（160字）
- **综合判断**：多空方向、价格区间、关键变量、逻辑链
- **操作策略**：做多/做空方向、建仓区间、止损位、止盈位、仓位建议、注意事项

【引用格式要求】⭐重要
- 所有事实性陈述必须使用上标[1][2][3]
- 技术指标数据直接引用，无需上标
- 每个维度都要有具体分析

⚠️ 再次强调：
1. 不要包含【行情回顾】（已在"一、市场走势回顾"中描述）
2. 【今日观点】和【操作建议】合并为【今日观点及操作建议】
3. 必须覆盖8大分析维度
4. 语言专业、逻辑严密、结论明确

请直接输出观点文本（按【早盘聚焦】【深度分析】【今日观点及操作建议】三部分结构输出）：
"""
        
        headers = {
            "Authorization": f"Bearer {DEEPSEEK_API_KEY}",
            "Content-Type": "application/json"
        }
        
        data = {
            "model": "deepseek-chat",
            "messages": [{"role": "user", "content": prompt}],
            "temperature": 0.2,  # 降低温度，提高准确性
            "max_tokens": 1200  # 增加token数，支持300-400字的8维度分析
        }
        
        response = requests.post(DEEPSEEK_API_URL, headers=headers, json=data, timeout=90)
        
        print(f"[DEBUG] AI生成主要观点 - 状态码: {response.status_code}")
        
        if response.status_code == 200:
            result = response.json()
            content = result['choices'][0]['message']['content'].strip()
            print(f"[DEBUG] AI生成主要观点 - 返回内容长度: {len(content)}字符")
            print(f"[DEBUG] AI生成主要观点 - 前100字符: {content[:100]}")
            return content
        else:
            error_msg = f"AI生成失败 (状态码: {response.status_code})"
            print(f"[ERROR] AI生成主要观点 - {error_msg}")
            if response.status_code == 401:
                print(f"[ERROR] API密钥无效或过期")
            elif response.status_code == 429:
                print(f"[ERROR] API调用频率过高或额度用完")
            return error_msg
            
    except Exception as e:
        error_msg = f"AI生成出错: {str(e)}"
        print(f"[ERROR] AI生成主要观点 - {error_msg}")
        import traceback
        print(f"[ERROR] 详细错误: {traceback.format_exc()}")
        return error_msg


# ============ 品种映射（用于输入提示）============
COMMODITY_EXAMPLES = {
    "铜": "CU2501",
    "铝": "AL2501",
    "锌": "ZN2501",
    "铅": "PB2501",
    "镍": "NI2501",
    "锡": "SN2501",
    "螺纹钢": "RB2501",
    "热卷": "HC2501",
    "铁矿石": "I2501",
    "焦炭": "J2501",
    "焦煤": "JM2501",
    "原油": "SC2501",
    "PTA": "TA2501",
    "甲醇": "MA2501",
    "豆粕": "M2501",
    "豆油": "Y2501",
    "玉米": "C2501",
    "白糖": "SR2501",
    "棉花": "CF2501",
    "黄金": "AU2512",
    "白银": "AG2512"
}


# 创建文件夹和文档保存路径
def create_folder_and_doc_path(custom_date):
    # 自动获取当前用户的桌面路径
    try:
        # 尝试获取桌面路径
        desktop = os.path.join(os.path.expanduser("~"), "Desktop")
        # 如果桌面不存在（中文系统可能是"桌面"），尝试中文路径
        if not os.path.exists(desktop):
            desktop = os.path.join(os.path.expanduser("~"), "桌面")
        # 如果还是不存在，就用用户主目录
        if not os.path.exists(desktop):
            desktop = os.path.expanduser("~")
    except:
        # 如果出错，使用当前目录
        desktop = "."

    base_path = os.path.join(desktop, "期货日报")
    folder_path = os.path.join(base_path, f"期货日报_{custom_date}")

    # 创建目录（如果不存在）
    try:
        if not os.path.exists(folder_path):
            os.makedirs(folder_path, exist_ok=True)
    except Exception as e:
        # 如果创建失败，使用当前目录
        folder_path = f"期货日报_{custom_date}"
        if not os.path.exists(folder_path):
            os.makedirs(folder_path, exist_ok=True)

    base_filename = "期货日报"
    filename = f"{base_filename}_{custom_date}.docx"
    doc_path = os.path.join(folder_path, filename)
    
    # 检查文件是否已存在且被占用，如果是则生成新文件名
    if os.path.exists(doc_path):
        counter = 1
        while True:
            new_filename = f"{base_filename}_{custom_date}_{counter}.docx"
            new_doc_path = os.path.join(folder_path, new_filename)
            if not os.path.exists(new_doc_path):
                doc_path = new_doc_path
                break
            # 尝试打开文件，如果可以打开说明没被占用，可以覆盖
            try:
                with open(new_doc_path, 'a'):
                    doc_path = new_doc_path
                    break
            except:
                counter += 1
                if counter > 10:  # 最多尝试10次
                    doc_path = new_doc_path
                    break
    
    return doc_path, folder_path


# 设置专业文档样式
def set_professional_doc_style(doc):
    """设置专业期货日报样式"""
    # 设置正文样式
    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")  # 使用宋体
    normal.font.size = Pt(12)
    
    # 设置段落格式
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.2


# 智能获取最近的交易日
def get_last_trading_day(symbol, report_date, max_days_back=7):
    """智能获取最近的交易日
    
    从report_date向前查找，找到最近一个有交易数据的日期
    避免周末、节假日导致的数据缺失问题
    
    Args:
        symbol: 合约代码
        report_date: 日报生成日期
        max_days_back: 最多向前查找的天数
    
    Returns:
        最近的交易日日期对象，如果找不到则返回None
    """
    st.info(f"🔍 正在查找最近的交易日...")
    
    for days_back in range(1, max_days_back + 1):
        check_date = report_date - timedelta(days=days_back)
        check_date_str = check_date.strftime('%Y-%m-%d')
        
        # 尝试获取该日期的数据
        try:
            start_time = check_date_str + ' 09:00:00'
            end_time = check_date_str + ' 15:00:00'
            
            df = ak.futures_zh_minute_sina(symbol=symbol, period="1")
            if not df.empty:
                df['datetime'] = pd.to_datetime(df['datetime'])
                day_data = df[(df['datetime'] >= start_time) & (df['datetime'] <= end_time)]
                
                if not day_data.empty:
                    if days_back > 1:
                        weekday = check_date.strftime('%A')
                        weekday_cn = {'Monday': '周一', 'Tuesday': '周二', 'Wednesday': '周三', 
                                    'Thursday': '周四', 'Friday': '周五', 'Saturday': '周六', 'Sunday': '周日'}
                        st.success(f"✅ 找到交易日：{check_date_str}（向前{days_back}天，{weekday_cn.get(weekday, weekday)}）")
                    else:
                        st.success(f"✅ 找到交易日：{check_date_str}")
                    return check_date
            
            time.sleep(0.1)  # 避免请求过快
        except Exception as e:
            continue
    
    st.warning(f"⚠️ 未找到有效交易日（已回溯{max_days_back}天）")
    return None


# 获取当天行情概述数据（智能查找最近交易日）
def get_market_trend_data(symbol, custom_date):
    """获取市场走势数据（包含白天盘和夜盘）
    
    对于日报生成日期（custom_date），获取最近交易日的完整交易数据：
    - 日盘：交易日 9:00-15:00
    - 夜盘：交易日 21:00 到次日凌晨（覆盖所有品种，最晚到02:30）
    
    Returns:
        day_description: 日盘描述
        night_description: 夜盘描述
        filtered_data: 过滤后的数据框
        market_data_dict: 市场数据字典
        market_date: 实际的交易日日期对象
    """
    try:
        # 日报生成日期
        report_date = custom_date
        
        # 智能查找最近的交易日（而不是简单的前一天）
        market_date = get_last_trading_day(symbol, report_date)
        
        if market_date is None:
            st.error("❌ 无法找到有效的交易日数据")
            return "", "", pd.DataFrame(), {}, None
        
        # 获取数据范围：从交易日的9:00到次日凌晨03:00（覆盖所有夜盘品种）
        next_day = market_date + timedelta(days=1)
        start_time = market_date.strftime('%Y-%m-%d') + ' 09:00:00'
        end_time = next_day.strftime('%Y-%m-%d') + ' 03:00:00'
        
        df = ak.futures_zh_minute_sina(symbol=symbol, period="1")
        
        if df.empty:
            st.error(f"⚠️ akshare未返回任何数据，可能原因：合约代码{symbol}不存在或格式错误")
            return "", "", pd.DataFrame(), {}, None
        
        df['datetime'] = pd.to_datetime(df['datetime'])
        filtered_data = df[(df['datetime'] >= start_time) & (df['datetime'] <= end_time)]
        
        if filtered_data.empty:
            st.error(f"⚠️ 指定时间范围内无数据")
            return "", "", pd.DataFrame(), {}, None

        # 日盘数据：交易日 9:00-15:00
        day_start = market_date.strftime('%Y-%m-%d') + ' 09:00:00'
        day_end = market_date.strftime('%Y-%m-%d') + ' 15:00:00'
        day_data = filtered_data[(filtered_data['datetime'] >= day_start) & (filtered_data['datetime'] <= day_end)]
        
        if day_data.empty:
            st.error(f"⚠️ 日盘时段（{day_start} 至 {day_end}）无数据")
            return "", "", pd.DataFrame(), {}, None
        
        # 获取开盘价和收盘价
        day_open_price = day_data.iloc[0]['open']
        day_close_price = day_data.iloc[-1]['close']

        high_price = day_data['high'].max()
        low_price = day_data['low'].min()
        price_change = day_close_price - day_open_price
        price_change_percentage = (price_change / day_open_price) * 100
        trend = "上涨" if price_change > 0 else "下跌" if price_change < 0 else "持平"
        
        day_description = (
            f"{market_date.strftime('%Y-%m-%d')}日{symbol}主力合约开盘价为{day_open_price}元/吨，最高价为{high_price}元/吨，"
            f"最低价为{low_price}元/吨，收盘价为{day_close_price}元/吨，较前一日{trend}了"
            f"{abs(price_change):.2f}元/吨，涨跌幅为{price_change_percentage:.2f}%。"
        )

        # 夜盘数据：交易日 21:00 到次日凌晨03:00（覆盖所有品种）
        night_start_time = market_date.strftime('%Y-%m-%d') + ' 21:00:00'
        night_end_time = next_day.strftime('%Y-%m-%d') + ' 03:00:00'
        night_data = df[(df['datetime'] >= night_start_time) & (df['datetime'] <= night_end_time)]
        
        night_description = ""
        market_data_dict = {
            'open': day_open_price,
            'close': day_close_price,
            'high': high_price,
            'low': low_price,
            'change': price_change,
            'change_pct': price_change_percentage
        }
        
        if night_data.empty:
            night_description = "夜盘数据不可用。"
            market_data_dict.update({
                'night_open': 'N/A',
                'night_close': 'N/A',
                'night_change': 'N/A',
                'night_change_pct': 'N/A'
            })
        else:
            night_open_price = night_data.iloc[0]['open']
            night_close_price = night_data.iloc[-1]['close']
            night_price_change = night_close_price - night_open_price
            night_price_change_percentage = (night_price_change / night_open_price) * 100
            night_trend = "上涨" if night_price_change > 0 else "下跌" if night_price_change < 0 else "持平"
            night_description = (
                f"夜盘走势：开盘价为{night_open_price}元/吨，收盘价为{night_close_price}元/吨，较开盘{night_trend}了"
                f"{abs(night_price_change):.2f}元/吨，涨跌幅为{night_price_change_percentage:.2f}%。"
            )
            market_data_dict.update({
                'night_open': night_open_price,
                'night_close': night_close_price,
                'night_change': night_price_change,
                'night_change_pct': night_price_change_percentage
            })

        # 返回日盘描述、夜盘描述、过滤后的数据框、市场数据字典、实际的交易日日期
        return day_description, night_description, filtered_data, market_data_dict, market_date
    except Exception as e:
        st.error(f"❌ 数据获取异常: {str(e)}")
        return f"获取市场走势数据失败: {e}", "", pd.DataFrame(), {}, None


# 创建K线图
def create_k_line_chart(data, symbol, folder_path):
    if data.empty:
        print("数据为空，无法生成K线图。")
        return None
    
    # 临时重置matplotlib配置，避免Linux字体错误
    import matplotlib
    matplotlib.rcParams['font.family'] = 'sans-serif'
    matplotlib.rcParams['font.sans-serif'] = ['DejaVu Sans']
    matplotlib.rcParams['axes.unicode_minus'] = False
    
    data.set_index('datetime', inplace=True)
    data = data[['open', 'high', 'low', 'close']]
    data.columns = ['Open', 'High', 'Low', 'Close']
    fig, ax = plt.subplots(figsize=(10, 6))
    mpf.plot(data, type='candle', style='charles', ax=ax)
    k_line_chart_path = os.path.join(folder_path, 'k_line_chart.png')
    plt.savefig(k_line_chart_path, dpi=100)
    plt.close(fig)
    return k_line_chart_path


# 获取新闻数据（使用增强搜索器）
def get_news_data_enhanced(commodity_name: str, serper_key: str = None, target_dates: list = None):
    """使用增强新闻搜索器获取新闻（支持指定日期列表，使用上标引用格式）"""
    try:
        searcher = EnhancedNewsSearcher()
        news_list = searcher.comprehensive_search(
            commodity=commodity_name,
            days_back=3,
            serper_key=serper_key,
            target_dates=target_dates
        )
        
        # 格式化新闻，显示所有新闻，使用上标引用格式[1][2][3]
        description = ""
        for i, news in enumerate(news_list, 1):
            # 优先使用content，如果没有就用title
            if news.get('content'):
                # 清理内容，去除多余空格、换行和省略号
                content = news['content'].strip().replace('...', '').replace('…', '')
                # 不限制长度，显示完整内容（如果太长Word会自动换行）
            else:
                # 如果没有content，就用title
                content = news['title']
            
            # 格式：序号. 内容[上标]
            description += f"{i}. {content}[{i}]\n\n"
        
        return description if description else "暂无相关新闻", news_list
    except Exception as e:
        return f"获取新闻数据失败: {e}", []


# 创建报告（专业版）
def create_report_professional(custom_date_str, symbol, commodity_name, user_description, main_view, 
                               user_news_content=None, serper_key=None):
    """创建专业期货日报
    
    Args:
        custom_date_str: 日期字符串
        symbol: 合约代码
        commodity_name: 品种名称
        user_description: 用户编辑的行情描述
        main_view: 用户编辑的主要观点
        user_news_content: 用户编辑的新闻资讯内容（可选）
        serper_key: Serper API密钥
    """
    custom_date = datetime.strptime(custom_date_str, '%Y-%m-%d')
    doc_path, folder_path = create_folder_and_doc_path(custom_date_str)
    
    # 获取市场数据（返回5个值：描述、夜盘描述、数据框、数据字典、实际交易日）
    market_trend_description, night_trend_description, market_data, market_data_dict, actual_market_date = get_market_trend_data(symbol=symbol, custom_date=custom_date)
    
    # 如果没有找到交易日，返回None
    if actual_market_date is None:
        st.error("无法生成报告，因为未找到有效的交易日数据。")
        return None
    
    # 生成资讯日期列表：从交易日到报告日期的所有日期
    news_date_list = []
    current_date = actual_market_date
    while current_date <= custom_date:
        news_date_list.append(current_date.strftime('%Y-%m-%d'))
        current_date += timedelta(days=1)
    
    st.info(f"📅 市场回顾：{actual_market_date.strftime('%Y-%m-%d')}  |  资讯搜索：{actual_market_date.strftime('%Y-%m-%d')} 至 {custom_date_str}（共{len(news_date_list)}天）")
    
    if market_data.empty:
        st.error("无法生成报告，因为市场数据为空。")
        return None
    
    # 使用用户编辑的新闻内容，如果没有则自动获取
    if user_news_content and user_news_content.strip():
        news_description = user_news_content
        # 从用户内容中提取引用，用于附录（简化处理，提取所有[数字]格式的引用）
        import re
        ref_numbers = re.findall(r'\[(\d+)\]', news_description)
        # 如果有引用，获取原始新闻列表以构建附录
        if ref_numbers and serper_key:
            _, news_list = get_news_data_enhanced(commodity_name, serper_key, news_date_list)
        else:
            news_list = []
    else:
        # 使用增强新闻搜索器（传入日期列表，保存新闻列表用于附录）
        news_description, news_list = get_news_data_enhanced(commodity_name, serper_key, news_date_list)
    
    k_line_chart_path = create_k_line_chart(market_data, symbol, folder_path)

    doc = Document()
    set_professional_doc_style(doc)

    # === 1. 添加标题 ===
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.add_run(f"期货日报")
    title_run.font.name = '黑体'
    title_run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
    title_run.font.size = Pt(18)
    title_run.font.bold = True
    title_run.font.color.rgb = RGBColor(31, 73, 125)

    # === 2. 添加品种和日期 ===
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_run = subtitle.add_run(f"{commodity_name} | {custom_date_str}")
    subtitle_run.font.size = Pt(14)
    subtitle_run.font.bold = True
    
    # 添加分隔线
    doc.add_paragraph("_" * 50)

    # === 3. 市场走势回顾 ===
    market_heading = doc.add_paragraph()
    market_heading_run = market_heading.add_run("一、市场走势回顾")
    market_heading_run.font.name = '黑体'
    market_heading_run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
    market_heading_run.font.size = Pt(14)
    market_heading_run.font.bold = True
    
    # 添加K线图
    if k_line_chart_path:
        doc.add_picture(k_line_chart_path, width=Inches(6))
        # 图片说明
        pic_caption = doc.add_paragraph()
        pic_caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
        pic_caption_run = pic_caption.add_run(f"图1：{commodity_name}日内走势图")
        pic_caption_run.font.size = Pt(10)
        pic_caption_run.italic = True
    
    # 行情描述
    market_content = doc.add_paragraph()
    market_content.add_run(user_description)

    # === 4. 主要观点 ===
    main_view_heading = doc.add_paragraph()
    main_view_heading_run = main_view_heading.add_run("二、主要观点")
    main_view_heading_run.font.name = '黑体'
    main_view_heading_run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
    main_view_heading_run.font.size = Pt(14)
    main_view_heading_run.font.bold = True
    
    main_view_content = doc.add_paragraph()
    main_view_content.add_run(main_view)

    # === 5. 市场资讯 ===
    news_heading = doc.add_paragraph()
    news_heading_run = news_heading.add_run("三、市场资讯")
    news_heading_run.font.name = '黑体'
    news_heading_run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
    news_heading_run.font.size = Pt(14)
    news_heading_run.font.bold = True
    
    news_content = doc.add_paragraph()
    news_content.add_run(news_description)

    # === 6. 附录：参考文献 ===
    doc.add_paragraph()  # 空行
    doc.add_paragraph("_" * 50)
    
    appendix_heading = doc.add_paragraph()
    appendix_heading_run = appendix_heading.add_run("附录：参考文献")
    appendix_heading_run.font.name = '黑体'
    appendix_heading_run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
    appendix_heading_run.font.size = Pt(12)
    appendix_heading_run.font.bold = True
    
    # 列出所有新闻引用（不限制数量，显示全部）
    for i, news in enumerate(news_list, 1):
        ref_para = doc.add_paragraph()
        ref_para.paragraph_format.left_indent = Inches(0.3)
        
        # 引用序号
        ref_run = ref_para.add_run(f"[{i}] ")
        ref_run.font.bold = True
        ref_run.font.size = Pt(10)
        
        # 新闻标题
        title_run = ref_para.add_run(news.get('title', '无标题'))
        title_run.font.size = Pt(10)
        
        # 来源和日期
        source_run = ref_para.add_run(f"\n    来源：{news.get('source', '未知')} | 日期：{news.get('date', 'N/A')}")
        source_run.font.size = Pt(9)
        source_run.italic = True
        
        # URL链接
        if news.get('url'):
            url_run = ref_para.add_run(f"\n    链接：{news.get('url', '无')}")
            url_run.font.size = Pt(9)
            url_run.font.color.rgb = RGBColor(0, 0, 255)  # 蓝色链接

    # === 7. 报告说明 ===
    doc.add_paragraph()  # 空行
    doc.add_paragraph("_" * 50)
    
    disclaimer = doc.add_paragraph()
    disclaimer_run = disclaimer.add_run("报告说明")
    disclaimer_run.font.size = Pt(10)
    disclaimer_run.font.bold = True
    
    disclaimer_content = doc.add_paragraph()
    disclaimer_content_run = disclaimer_content.add_run(
        "本报告基于公开信息和市场数据编制，仅供参考。期货市场存在风险，投资需谨慎。"
        "本报告不构成投资建议，投资者应根据自身情况独立决策并承担风险。"
    )
    disclaimer_content_run.font.size = Pt(9)
    disclaimer_content_run.italic = True
    
    # === 8. 页脚信息 ===
    footer_para = doc.add_paragraph()
    footer_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    footer_run = footer_para.add_run(f"{custom_date_str}")
    footer_run.font.size = Pt(9)
    footer_run.italic = True

    # 保存文档（带重试机制）
    max_retries = 3
    for attempt in range(max_retries):
        try:
            doc.save(doc_path)
            return doc_path
        except PermissionError as e:
            if attempt < max_retries - 1:
                # 如果不是最后一次尝试，生成新的文件名
                import time
                time.sleep(0.5)  # 等待0.5秒
                base_path = os.path.dirname(doc_path)
                base_name = os.path.basename(doc_path).replace('.docx', '')
                new_doc_path = os.path.join(base_path, f"{base_name}_副本{attempt+1}.docx")
                doc_path = new_doc_path
            else:
                # 最后一次尝试失败，抛出错误并给出提示
                st.error(f"❌ 保存失败：文件可能正在被使用")
                st.warning(f"💡 解决方法：\n1. 关闭已打开的Word文档\n2. 重新点击生成按钮\n3. 或等待文件自动保存为副本")
                raise
        except Exception as e:
            st.error(f"❌ 保存失败：{str(e)}")
            raise
    
    return doc_path


# ============ Streamlit应用 ============
st.set_page_config(page_title="期货日报生成器（AI赋能版）", page_icon="📊", layout="wide")

# ============ 侧边栏：系统信息 ============
st.sidebar.title("📊 期货日报生成器")
st.sidebar.markdown("**AI赋能版**")
st.sidebar.markdown("---")

# 系统状态
st.sidebar.subheader("✅ 系统状态")
st.sidebar.success("🤖 DeepSeek AI - 已就绪")
st.sidebar.success("🔍 Serper搜索 - 已就绪")
st.sidebar.success("📈 数据接口 - 已就绪")

st.sidebar.markdown("---")

# 功能说明
st.sidebar.subheader("🎯 核心功能")
st.sidebar.markdown("""
- 📊 自动生成K线图
- 🤖 AI智能行情分析
- 🧠 8维度专业观点
- 📰 多源新闻聚合
- 📄 一键生成日报
""")

st.sidebar.markdown("---")
st.sidebar.info("💡 所有AI功能已内置配置，可直接使用")

# ============ 主界面 ============
st.title("📊 期货日报生成器（AI赋能版）")
st.write("**created by 7haoge (953534947@qq.com)**")

st.markdown("---")

# 显示系统说明
with st.expander("📖 系统使用说明", expanded=False):
    st.markdown("""
    ### 📊 系统简介
    
    **期货日报生成器（AI赋能版）** 是一款专业的期货日报自动生成工具，利用AI技术帮助您快速生成专业的期货市场日报。
    
    ### ✨ 核心功能
    
    1. **📈 K线图自动生成**
       - 获取期货品种的实时行情数据
       - 自动绘制专业K线走势图
       - 包含日盘和夜盘完整数据
    
    2. **🤖 AI智能分析**
       - AI自动生成行情描述
       - AI综合分析生成投资观点
       - 基于DeepSeek大模型驱动
    
    3. **📰 多源新闻聚合**
       - 自动抓取东方财富、金融界等财经网站新闻
       - 智能筛选与品种相关的资讯
       - 按相关性排序，精选最重要的10条
    
    4. **📄 专业报告输出**
       - 符合机构标准的Word文档格式
       - 包含完整的报告结构和免责声明
       - 一键下载，直接可用
    
    ### 🚀 使用流程
    
    **第一步：输入基本信息**
    - 选择日期
    - 输入品种名称（如：铜、螺纹钢、PTA等）
    - 输入合约代码（如：CU2501、RB2501等）
    
    **第二步：生成K线图**
    - 点击"生成K线图"按钮
    - 系统自动获取市场数据并生成图表
    - 自动获取相关新闻资讯
    
    **第三步：编辑报告内容（AI智能 + 人工审核）**
    
    - **📝 行情描述**：
      - 点击"AI生成行情描述"：基于真实行情数据，自动生成专业市场描述
      - 或手动编辑，灵活调整
    
    - **💡 主要观点（专业版）**：⭐ 基于8大专业维度系统分析
      - 点击"AI生成主要观点（专业版）"
      - AI将综合以下8个维度进行专业分析：
        1. 📊 技术面分析（MA、MACD、RSI、布林带）
        2. 📦 基本面分析（库存、仓单、供需、基差）
        3. 💰 资金面分析（持仓席位、主力动向）
        4. 🔗 产业链分析（上下游价格、利润）
        5. 📜 政策面分析（国家政策、行业政策）
        6. 🌍 进出口分析（海关数据、贸易情况）
        7. 📰 市场情绪（新闻舆情、投资者情绪）
        8. ⚠️ 风险因素（主要风险、不确定性）
      - 生成后可手动编辑和调整
    
    - **📰 新闻资讯（NEW!）**：
      - 点击"AI生成新闻资讯"：AI自动整理指定日期搜索到的所有新闻
      - 包含：商品新闻、库存数据、基差数据、持仓变化等多维度信息
      - 只展示新闻主体内容，自动标注引用[1][2][3]
      - 自动去重并统一格式
      - 支持手动编辑、删减或补充
    
    **第四步：生成完整日报**
    - 点击"生成完整日报"按钮
    - Word文档自动保存到桌面「期货日报」文件夹
    - 也可以直接下载
    
    ### 💡 使用技巧
    
    - **合约代码格式**：品种代码（大写）+ 年份（2位）+ 月份（2位）
      - 例如：CU2501（2025年1月交割的铜合约）
    - **AI生成内容**：可作为初稿参考，建议根据实际情况调整
    - **日报生成时间**：建议在每日收盘后15:30-16:00生成
    
    ### ⚠️ 注意事项
    
    - 系统需要稳定的网络连接获取数据
    - 合约代码必须准确，否则无法获取数据
    - 生成的报告仅供参考，不构成投资建议
    - 日报将保存在桌面的"期货日报"文件夹中
    
    ### 📞 技术支持
    
    如有问题或建议，欢迎联系：
    - **作者**：7haoge
    - **邮箱**：953534947@qq.com
    """)

# 主界面
col1, col2 = st.columns([1, 1])

with col1:
    st.subheader("📅 基本信息")
    custom_date = st.date_input("请选择日期", datetime.now())
    
    # 品种名输入
    commodity_name = st.text_input(
        "请输入品种名称",
        placeholder="例如：铜、螺纹钢、PTA、豆粕、白银等",
        help="直接输入品种中文名称"
    )
    
    # 显示合约代码示例
    example_contract = ""
    if commodity_name and commodity_name in COMMODITY_EXAMPLES:
        example_contract = COMMODITY_EXAMPLES[commodity_name]
    
    # 合约代码输入
    full_contract = st.text_input(
        "请输入完整品种合约",
        placeholder=f"例如：{example_contract if example_contract else 'CU2501、RB2501、AG2512'}",
        help="格式：品种代码+年月，如CU2501表示2025年1月交割的铜合约"
    )

with col2:
    pass  # 数据预览部分已删除

st.markdown("---")

# 初始化session state
if 'market_data_dict' not in st.session_state:
    st.session_state.market_data_dict = {}
if 'news_list' not in st.session_state:
    st.session_state.news_list = []
if 'day_description' not in st.session_state:
    st.session_state.day_description = ""
if 'night_description' not in st.session_state:
    st.session_state.night_description = ""
if 'ai_generated_description' not in st.session_state:
    st.session_state.ai_generated_description = ""
if 'ai_generated_view' not in st.session_state:
    st.session_state.ai_generated_view = ""
if 'ai_generated_news' not in st.session_state:
    st.session_state.ai_generated_news = ""
if 'professional_data' not in st.session_state:
    st.session_state.professional_data = {}
if 'commodity_name' not in st.session_state:
    st.session_state.commodity_name = ""
if 'full_contract' not in st.session_state:
    st.session_state.full_contract = ""
if 'custom_date' not in st.session_state:
    st.session_state.custom_date = datetime.now()
if 'temp_ai_desc' not in st.session_state:
    st.session_state.temp_ai_desc = ""
if 'temp_ai_view' not in st.session_state:
    st.session_state.temp_ai_view = ""
if 'temp_ai_news' not in st.session_state:
    st.session_state.temp_ai_news = ""

# K线图生成
if st.button("🎨 生成K线图", type="primary"):
    if not full_contract or not commodity_name:
        st.error("❌ 请先输入品种名称和完整合约代码")
    else:
        # 保存用户输入到session state
        st.session_state.commodity_name = commodity_name
        st.session_state.full_contract = full_contract
        st.session_state.custom_date = custom_date
        
        with st.spinner("正在生成K线图..."):
            custom_date_str = custom_date.strftime('%Y-%m-%d')
            
            # 获取市场数据（返回5个值）
            day_description, night_description, market_data, market_data_dict, actual_market_date = get_market_trend_data(full_contract, custom_date)
            
            # 如果找不到交易日，停止
            if actual_market_date is None:
                st.error("无法找到有效的交易日数据，请检查合约代码或选择其他日期")
                st.stop()
            
            # 保存到session state
            st.session_state.market_data_dict = market_data_dict
            st.session_state.actual_market_date = actual_market_date  # 保存实际交易日
            st.session_state.market_data_df = market_data  # 保存DataFrame用于技术指标计算
            st.session_state.day_description = day_description
            st.session_state.night_description = night_description
            
            # 生成资讯日期列表：从交易日到报告日期的所有日期
            news_date_list = []
            current_date = actual_market_date
            while current_date <= custom_date:
                news_date_list.append(current_date.strftime('%Y-%m-%d'))
                current_date += timedelta(days=1)
            
            st.info(f"📰 资讯搜索范围：{actual_market_date.strftime('%Y-%m-%d')} 至 {custom_date_str}（共{len(news_date_list)}天）")
            
            # 获取新闻（传入日期列表）
            news_description, news_list = get_news_data_enhanced(
                commodity_name, 
                SERPER_API_KEY,
                target_dates=news_date_list  # 传入日期列表
            )
            st.session_state.news_list = news_list
            
            if not market_data.empty:
                k_line_chart_path = create_k_line_chart(market_data, full_contract, ".")
                
                if k_line_chart_path:
                    col_img1, col_img2 = st.columns([2, 1])
                    with col_img1:
                        st.image(k_line_chart_path, caption=f"{commodity_name}昨日K线图", use_container_width=True)
                    with col_img2:
                        st.write("**📊 昨日走势：**")
                        st.write(day_description)
                        st.write("")
                        st.write("**🌙 夜盘走势：**")
                        st.write(night_description)
                else:
                    st.error("❌ 无法生成K线图")
            else:
                st.error("❌ 无法获取市场数据，请检查合约代码是否正确")

st.markdown("---")

# 内容编辑区域
st.subheader("✍️ 编辑报告内容")

# 行情描述区域
st.markdown("### 📝 行情描述")
col_desc1, col_desc2 = st.columns([3, 1])

with col_desc1:
    # 确定显示的内容：优先显示AI生成的，否则显示自动生成的
    default_description = st.session_state.get('ai_generated_description', '')
    if not default_description:
        default_description = st.session_state.get('day_description', '') + '\n\n' + st.session_state.get('night_description', '')
    
    user_description = st.text_area(
        "请输入行情描述（可采用自动生成的文案或自行编辑，也可以使用AI生成）",
        value=default_description,
        height=200,
        key="user_description"
    )

with col_desc2:
    st.write("")
    st.write("")
    
    if st.button("🤖 AI生成行情描述", use_container_width=True, key="btn_gen_desc"):
        if not st.session_state.get('market_data_dict'):
            st.warning("⚠️ 请先生成K线图以获取市场数据")
        elif not st.session_state.get('commodity_name'):
            st.warning("⚠️ 请先输入品种名称并生成K线图")
        else:
            try:
                with st.spinner("🤖 AI正在生成..."):
                    # 获取技术指标
                    technical_indicators = {}
                    if 'market_data_df' in st.session_state and not st.session_state.market_data_df.empty:
                        technical_indicators = calculate_technical_indicators(st.session_state.market_data_df)
                    
                    ai_desc = ai_generate_market_description(
                        st.session_state.market_data_dict,
                        st.session_state.commodity_name,
                        st.session_state.custom_date.strftime('%Y-%m-%d'),
                        technical_indicators=technical_indicators,
                        market_date=st.session_state.get('actual_market_date')  # 传递实际交易日
                    )
                
                if ai_desc and len(ai_desc) > 50:
                    # 保存到session_state
                    st.session_state.temp_ai_desc = ai_desc
                    st.success("✅ 生成成功！内容显示在下方")
                else:
                    st.error(f"❌ 生成失败: {ai_desc}")
            except Exception as e:
                st.error(f"❌ 异常: {str(e)}")
    
    # 显示已生成的内容（如果有）
    if st.session_state.get('temp_ai_desc'):
        st.write("📝 生成的行情描述：")
        st.text_area("AI生成的行情描述", value=st.session_state.temp_ai_desc, height=200, key="display_ai_desc", label_visibility="collapsed")
        st.caption("💡 请复制上面的内容到上方输入框")

# 主要观点区域
st.markdown("### 💡 主要观点")
col_view1, col_view2 = st.columns([3, 1])

with col_view1:
    # 确定显示的内容：优先显示AI生成的
    default_view = st.session_state.get('ai_generated_view', '')
    
    main_view = st.text_area(
        "请输入主要观点（可自行编辑或AI生成）",
        value=default_view,
        height=200,
        key="main_view",
        placeholder="输入您对市场的主要判断和投资建议..."
    )

with col_view2:
    st.write("")
    st.write("")
    if st.button("🧠 AI生成主要观点（专业版）", use_container_width=True, key="btn_gen_view"):
        if not st.session_state.get('market_data_dict'):
            st.warning("⚠️ 请先生成K线图以获取市场数据")
        elif not st.session_state.get('commodity_name'):
            st.warning("⚠️ 请先输入品种名称并生成K线图")
        else:
            try:
                # 第1步：计算技术指标
                with st.spinner("📊 正在计算技术指标..."):
                    market_data_df = st.session_state.get('market_data_df', pd.DataFrame())
                    technical_indicators = calculate_technical_indicators(market_data_df)
                
                # 第2步：搜索8大维度专业数据
                with st.spinner("🔍 正在搜索8大维度专业数据..."):
                    searcher = EnhancedNewsSearcher()
                    professional_data = searcher.search_professional_data(
                        st.session_state.commodity_name,
                        SERPER_API_KEY,
                        st.session_state.custom_date.strftime('%Y-%m-%d')
                    )
                    st.session_state.professional_data = professional_data if professional_data else {}
                
                # 第3步：AI综合分析生成观点
                with st.spinner("🤖 AI正在进行8大维度专业分析..."):
                    ai_view = ai_generate_main_view(
                        st.session_state.commodity_name,
                        st.session_state.custom_date.strftime('%Y-%m-%d'),
                        st.session_state.market_data_dict,
                        st.session_state.news_list,
                        professional_data,
                        technical_indicators,
                        market_date=st.session_state.get('actual_market_date')  # 传递实际交易日
                    )
                
                if ai_view and len(ai_view) > 50:
                    # 保存到session_state
                    st.session_state.temp_ai_view = ai_view
                    st.success("✅ 主要观点生成完成！内容显示在下方")
                else:
                    st.error(f"❌ 生成失败: {ai_view}")
            except Exception as e:
                st.error(f"❌ 异常: {str(e)}")
    
    # 显示已生成的内容（如果有）
    if st.session_state.get('temp_ai_view'):
        st.write("📝 生成的主要观点：")
        st.text_area("AI生成的主要观点", value=st.session_state.temp_ai_view, height=250, key="display_ai_view", label_visibility="collapsed")
        st.caption("💡 请复制上面的内容到上方输入框")

# 新闻资讯区域
st.markdown("### 📰 新闻资讯")
col_news1, col_news2 = st.columns([3, 1])

with col_news1:
    # 确定显示的内容：优先显示AI生成的
    default_news = st.session_state.get('ai_generated_news', '')
    
    news_content = st.text_area(
        "请输入新闻资讯（可自行编辑或AI生成）",
        value=default_news,
        height=300,
        key="news_content",
        placeholder="""输入或生成新闻资讯，格式示例：

1. 今日生猪期货主力合约收于11325元，日内下跌2.12%，创近期新低[1]

2. 全国生猪均价跌至11.19元/公斤，环比下跌0.26元，月环比下跌16.04%[2]

3. 交易所仓单数据显示，生猪仓单环比增加3000吨至15000吨，供应压力持续[3]

...

💡 提示：可手动输入，也可点击右侧"AI生成"按钮自动整理"""
    )

with col_news2:
    st.write("")
    st.write("")
    if st.button("📰 AI生成新闻资讯", use_container_width=True, key="btn_gen_news"):
        if not st.session_state.get('news_list'):
            st.warning("⚠️ 请先生成K线图以获取新闻数据")
        elif not st.session_state.get('commodity_name'):
            st.warning("⚠️ 请先输入品种名称并生成K线图")
        else:
            try:
                # 如果professional_data不存在，先搜索一次
                professional_data = st.session_state.get('professional_data')
                if not professional_data:
                    with st.spinner("🔍 正在搜索专业数据..."):
                        searcher = EnhancedNewsSearcher()
                        professional_data = searcher.search_professional_data(
                            st.session_state.commodity_name,
                            SERPER_API_KEY,
                            st.session_state.custom_date.strftime('%Y-%m-%d')
                        )
                        st.session_state.professional_data = professional_data if professional_data else {}
                
                with st.spinner("🤖 AI正在整理新闻资讯..."):
                    ai_news = ai_generate_news_summary(
                        st.session_state.commodity_name,
                        st.session_state.custom_date.strftime('%Y-%m-%d'),
                        st.session_state.news_list,
                        professional_data or {}
                    )
                
                if ai_news and len(ai_news) > 50:
                    # 保存到session_state
                    st.session_state.temp_ai_news = ai_news
                    st.success("✅ 新闻资讯生成完成！内容显示在下方")
                else:
                    st.error(f"❌ 生成失败: {ai_news}")
            except Exception as e:
                st.error(f"❌ 异常: {str(e)}")
    
    # 显示已生成的内容（如果有）
    if st.session_state.get('temp_ai_news'):
        st.write("📝 生成的新闻资讯：")
        st.text_area("AI生成的新闻资讯", value=st.session_state.temp_ai_news, height=300, key="display_ai_news", label_visibility="collapsed")
        st.caption("💡 请复制上面的内容到上方输入框")
    else:
        st.write("")
        st.info("💡 AI会整理所有搜索到的新闻")

st.markdown("---")

# 生成日报按钮
if st.button("📄 生成完整日报", type="primary"):
    if not full_contract or not commodity_name:
        st.error("❌ 请先输入品种名称和合约代码")
    elif not user_description or not main_view:
        st.error("❌ 请填写行情描述和主要观点")
    elif not news_content:
        st.warning("⚠️ 建议填写新闻资讯内容，或点击AI生成")
        st.info("💡 您也可以留空新闻资讯，继续生成日报")
    
    # 无论新闻资讯是否填写，都允许生成（兼容性）
    if (full_contract and commodity_name and user_description and main_view):
        try:
            with st.spinner("正在生成日报，请稍候..."):
                custom_date_str = custom_date.strftime('%Y-%m-%d')
                doc_path = create_report_professional(
                    custom_date_str,
                    full_contract,
                    commodity_name,
                    user_description,
                    main_view,
                    user_news_content=news_content,  # 传入用户编辑的新闻内容
                    serper_key=SERPER_API_KEY
                )
                
                if doc_path:
                    st.success("✅ 日报生成成功！")
                    
                    # 显示保存路径和文件名
                    st.info(f"📁 报告已保存至：`{os.path.dirname(doc_path)}`")
                    st.info(f"📄 文件名：`{os.path.basename(doc_path)}`")
                    
                    # 提示如何查看
                    st.markdown("💡 **查看报告：** 可在桌面「期货日报」文件夹中找到生成的Word文档")
                    
                    with open(doc_path, "rb") as f:
                        st.download_button(
                            label="📥 下载日报",
                            data=f,
                            file_name=os.path.basename(doc_path),
                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                            use_container_width=True
                        )
                    
                    st.balloons()
                else:
                    st.error("❌ 日报生成失败，请重试")
        except PermissionError:
            st.error("❌ 文件保存失败：权限被拒绝")
            st.warning("""
            **💡 常见原因和解决方法：**
            
            1. **Word文档正在打开** 
               - 请关闭桌面「期货日报」文件夹中已打开的同名Word文档
               - 然后重新点击"生成完整日报"按钮
            
            2. **文件被占用**
               - 系统已自动尝试生成带序号的副本文件
               - 请查看文件夹中是否有「期货日报_日期_副本1.docx」等文件
            
            3. **权限不足**
               - 确保您有桌面文件夹的写入权限
               - 尝试以管理员身份运行程序
            """)
        except Exception as e:
            st.error(f"❌ 生成失败：{str(e)}")
            st.warning("""
            **💡 请检查：**
            - 网络连接是否正常
            - 所有必填信息是否完整
            - 是否有足够的磁盘空间
            - API密钥是否正确配置
            """)

# 页脚
st.markdown("---")
st.markdown("""
<div style='text-align: center; color: gray; padding: 20px;'>
    <p>📊 期货日报生成器（AI赋能版）</p>
    <p>专业 | 智能 | 效率</p>
    <p>Powered by DeepSeek</p>
    <p>Created by 7haoge（953534947@qq.com）</p>
</div>
""", unsafe_allow_html=True)

